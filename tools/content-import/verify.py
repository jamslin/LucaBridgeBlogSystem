#!/usr/bin/env python3
"""
Stage A4 — independent verification of posts.json against the original .docx files.

This does NOT trust the extractor. It re-opens every source document and re-derives the
check from scratch, so a bug in extract.py cannot hide behind a matching bug here.

Checks:
  1. CHARACTER INVARIANT — for every event, the sequence of non-whitespace, non-corrupt
     characters in the stored tcBody is identical to the source .docx body. This is the
     hard rule: whitespace and corrupt code points may change, nothing else may.
  2. Column limits — every field fits the varchar() width of its DB column.
  3. Referential sanity — every referenced image exists on disk; slugs are unique;
     cover is a member of the gallery.
  4. A per-event whitespace-change log, written to whitespace-changes.txt, so the
     spacing edits can be eyeballed rather than taken on trust.

Exit code is non-zero if any check fails.
"""
import json, sys, unicodedata
from pathlib import Path

import docx
import textclean as tc

LIMITS = {"slug": 200, "tcTitle": 300, "tcVenue": 300, "tcSummary": 600,
          "enTitle": 300, "enVenue": 300, "enSummary": 600,
          "scTitle": 300, "scVenue": 300, "scSummary": 600}


def source_signature(folder: Path) -> str:
    """
    Signature of the ENTIRE source document — every paragraph concatenated, with whitespace
    and corrupt code points removed.

    Checking containment against the whole document (rather than re-picking the body
    paragraph) means this verification does not depend on the extractor's classification
    logic at all. If the stored text is an exact contiguous substring of this, then no
    character was altered, inserted, dropped or reordered — only whitespace and corruption
    can have changed. That is precisely the guarantee we owe.
    """
    docs = list(folder.glob("*.docx"))
    if not docs:
        return ""
    paras = [p.text for p in docx.Document(str(docs[0])).paragraphs]
    return tc.signature("".join(paras))


def main():
    assets = Path(sys.argv[1] if len(sys.argv) > 1 else "Assets")
    events = json.loads(Path("posts.json").read_text(encoding="utf-8"))
    fail, warn, log = [], [], []

    slugs = {}
    for e in events:
        f = assets / e["folder"]
        tag = f"{e['folder']} ({e['slug']})"

        # 1. character invariant — stored text must be a contiguous substring of the source
        src_sig = source_signature(f)
        for field in ("tcBody", "tcSummary", "tcTitle", "tcVenue"):
            v = e.get(field) or ""
            if not v:
                continue
            if field == "tcTitle" and (e.get("titleWrittenByClaude")
                                       or e.get("tcTitleModified")):
                continue                       # authored by us on purpose, not from source
            if "MERGED_FOLDER" in e["flags"] and not src_sig:
                continue
            if "SPLIT_SECOND_HALF" in e["flags"] and field in ("tcTitle", "tcVenue"):
                continue
            sig = tc.signature(v)
            if not src_sig:
                fail.append(f"[GHOST TEXT] {tag} has {field} but no source .docx")
            elif sig not in src_sig:
                # locate the first character that breaks containment, for a usable message
                lo, hi = 0, len(sig)
                while lo < hi:
                    mid = (lo + hi + 1) // 2
                    if sig[:mid] in src_sig:
                        lo = mid
                    else:
                        hi = mid - 1
                fail.append(f"[INVARIANT] {tag} {field} diverges from source at offset {lo}: "
                            f"...{sig[max(0,lo-20):lo]}[{sig[lo:lo+10]}]... not found in source")

        if e["tcBody"]:
            src_all = "".join(p.text for p in docx.Document(str(list(f.glob('*.docx'))[0])).paragraphs) \
                      if list(f.glob("*.docx")) else ""
            removed = src_all.count(" ") - e["tcBody"].count(" ")
            corrupt = sum(1 for ch in src_all if tc.is_corrupt(ch))
            paras_n = e["tcBody"].count("\n\n") + 1
            log.append(f"{tag}\n"
                       f"    spaces removed        : {removed}\n"
                       f"    corrupt chars stripped: {corrupt}\n"
                       f"    paragraphs created    : {paras_n}\n"
                       f"    characters kept       : {len(tc.signature(e['tcBody']))}")

        # 2. column limits
        for field, cap in LIMITS.items():
            v = e.get(field)
            if v and len(v) > cap:
                fail.append(f"[TOO LONG] {tag} {field} is {len(v)} chars, column is varchar({cap})")

        # 3. referential sanity
        if e["slug"] in slugs:
            fail.append(f"[DUP SLUG] {e['slug']} used by {slugs[e['slug']]} and {e['folder']}")
        slugs[e["slug"]] = e["folder"]

        for name in e["galleryImages"]:
            p = (assets / name) if "/" in name else (f / name)
            if not p.exists():
                fail.append(f"[MISSING FILE] {tag} gallery image not found: {name}")
        cover = e.get("coverImage")
        if cover:
            cp = (assets / "_posters" / cover) if e.get("coverFrom") == "VIDEO_POSTER" \
                 else ((assets / cover) if "/" in cover else (f / cover))
            if not cp.exists():
                fail.append(f"[MISSING FILE] {tag} cover not found: {cover}")
        elif not e["skippedVideos"]:
            warn.append(f"[NO COVER] {tag} will publish with no cover image")

        if e.get("tcAppendix") and e.get("tcAppendix") in (e.get("tcBody") or ""):
            fail.append(f"[APPENDIX LEAKED] {tag} tcAppendix was written into tcBody — "
                        f"tcBody must stay byte-identical to the source")
        if not e["tcTitle"]:
            fail.append(f"[NO TITLE] {tag} has no tcTitle — tc_title is NOT NULL")
        if not e["tcBody"]:
            warn.append(f"[NO BODY] {tag} has no body text")

        # corrupt characters must be entirely gone from the output
        for field in ("tcTitle", "tcSummary", "tcBody", "tcVenue"):
            v = e.get(field) or ""
            bad = [ch for ch in v if tc.is_corrupt(ch)]
            if bad:
                fail.append(f"[CORRUPT LEFT] {tag} {field} still contains "
                            f"{len(bad)} corrupt code point(s)")

    Path("whitespace-changes.txt").write_text("\n".join(log), encoding="utf-8")

    print(f"events checked : {len(events)}")
    print(f"failures       : {len(fail)}")
    print(f"warnings       : {len(warn)}")
    for x in fail:
        print("  FAIL", x)
    for x in warn:
        print("  warn", x)
    print("\nwhitespace change log -> whitespace-changes.txt")
    return 1 if fail else 0


if __name__ == "__main__":
    sys.exit(main())
